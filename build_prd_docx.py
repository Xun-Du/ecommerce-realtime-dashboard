from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "实时电商AB测试与归因分析看板_PRD.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(23, 43, 77)
MUTED = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
SUCCESS = RGBColor(36, 99, 57)
RISK = RGBColor(155, 28, 28)
GOLD = RGBColor(122, 90, 0)


def set_run_font(run, name="Calibri", east_asia="STHeiti", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    set_paragraph_spacing(p, before=0, after=0, line=1.08)
    r = p.add_run(str(text))
    set_run_font(r, size=size, color=color or INK, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], LIGHT_GRAY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, size=font_size, align=align)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_paragraph_spacing(p, before=16 if level == 1 else 10, after=6 if level <= 2 else 4, line=1.1)
    for run in p.runs:
        set_run_font(
            run,
            size={1: 16, 2: 13, 3: 12}.get(level, 11),
            color=BLUE if level <= 2 else DARK_BLUE,
            bold=True,
        )
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.12)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.8, color=INK, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.8, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.8, color=INK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph_spacing(p, before=0, after=4, line=1.12)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, before=0, after=4, line=1.12)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    return p


def add_callout(doc, title, body, color="E8EEF5"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    shade_cell(cell, color)
    set_cell_margins(cell, top=140, start=160, bottom=140, end=160)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2, line=1.1)
    r = p.add_run(title)
    set_run_font(r, size=10.8, color=INK, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0, line=1.12)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.3, color=INK)
    doc.add_paragraph()


def add_metadata(doc, pairs):
    for label, value in pairs:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=2, line=1.08)
        r1 = p.add_run(f"{label}: ")
        set_run_font(r1, size=10.5, color=INK, bold=True)
        r2 = p.add_run(value)
        set_run_font(r2, size=10.5, color=INK)


def add_rule(doc):
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(8)
    p_format.space_after = Pt(12)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "STHeiti")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0, line=1)
    rh = header.add_run("实时电商 A/B 测试与归因分析看板 PRD")
    set_run_font(rh, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rf = footer.add_run("Data Product Portfolio | Page ")
    set_run_font(rf, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = footer.add_run()
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=MUTED)
    return doc


def build_doc():
    doc = setup_document()

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=4, line=1.05)
    r = p.add_run("产品需求文档 PRD")
    set_run_font(r, size=12, color=MUTED, bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=6, line=1.05)
    tr = title.add_run("实时电商 A/B 测试与归因分析看板")
    set_run_font(tr, size=24, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=14, line=1.1)
    sr = subtitle.add_run("面向增长实验决策、渠道归因和实时经营监控的数据产品经理作品集项目")
    set_run_font(sr, size=12.5, color=MUTED)

    add_metadata(
        doc,
        [
            ("文档类型", "Product Requirements Document"),
            ("项目定位", "数据产品经理作品集项目"),
            ("目标岗位方向", "增长/实验方向数据产品经理"),
            ("业务场景", "综合电商平台"),
            ("版本", "V1.0"),
            ("日期", "2026-07-29"),
            ("适用读者", "增长产品、数据产品、数据分析、后端研发、前端研发、运营投放团队"),
        ],
    )
    add_rule(doc)
    add_callout(
        doc,
        "核心产品判断",
        "本产品不是单纯的数据展示看板，而是围绕“发现问题 - 定位原因 - 验证策略 - 做出上线决策”的增长实验闭环设计。PRD 重点体现数据产品经理对业务目标、指标口径、实验结论和跨团队交付边界的定义能力。",
    )

    add_heading(doc, "1. 文档信息", 1)
    add_table(
        doc,
        ["字段", "说明"],
        [
            ("项目名称", "实时电商 A/B 测试与归因分析看板"),
            ("项目阶段", "MVP / 作品集演示版"),
            ("核心目标", "帮助综合电商增长团队实时评估实验效果、定位漏斗流失、理解渠道贡献，并形成可执行的策略决策。"),
            ("产品边界", "第一版不包含登录权限、多租户、真实支付、真实广告平台回传和复杂实时流计算。"),
            ("交付形态", "Supabase Postgres + FastAPI + Streamlit + Plotly 的端到端可访问 Demo。"),
        ],
        [1.45, 5.05],
    )

    add_heading(doc, "2. 项目背景", 1)
    add_body(
        doc,
        "综合电商平台在活动页改版、推荐策略优化、结算流程简化和渠道投放调整中，经常需要通过 A/B 测试判断新策略是否真正提升转化。传统分析流程通常依赖离线报表和人工取数，实验结论滞后，且不同团队对转化率、GMV、渠道贡献的口径理解不一致，容易导致决策分歧。"
    )
    add_body(
        doc,
        "本项目通过模拟实时事件流、后端指标计算接口和交互式 Dashboard，把增长实验分析链路产品化：业务方可以看到实时经营表现，数据产品经理可以维护统一指标口径，增长 PM 可以基于显著性检验判断实验是否具备上线价值。"
    )
    add_callout(
        doc,
        "业务痛点",
        "离线报表响应慢、实验判断依赖人工、漏斗流失定位不直观、渠道归因口径不透明、预警阈值无法由业务侧灵活配置。",
        color="FFF4CE",
    )

    add_heading(doc, "3. 产品目标", 1)
    for item in [
        "实时监控 DAU、GMV、购买转化率、平均订单金额等核心经营指标。",
        "自动评估 A/B 实验效果，输出业务可读的实验结论。",
        "通过 click → add_to_cart → buy 漏斗定位转化流失环节。",
        "支持首次触达和最后触达归因模型切换，辅助判断渠道贡献。",
        "支持实验分流比例和指标预警阈值配置，形成策略管理闭环。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4. 目标用户与使用场景", 1)
    add_table(
        doc,
        ["用户角色", "核心诉求", "典型使用路径"],
        [
            ("增长产品经理", "判断实验策略是否值得继续、停止或全量上线。", "查看实验评估模块，比较 A/B 转化率、GMV、uplift 和 p-value，读取自动结论。"),
            ("数据产品经理", "统一指标口径，沉淀实验分析能力，推动数据链路产品化。", "维护 PRD 口径、验收 API 指标结果、定义异常与阈值规则。"),
            ("运营/投放同学", "理解不同渠道的订单和 GMV 贡献，调整预算或活动策略。", "切换归因模型，查看渠道贡献排行，并结合漏斗流失判断渠道质量。"),
            ("数据分析师", "复核实验结论，定位异常波动来源。", "按时间、实验组、渠道下钻，检查样本量、漏斗转化和指标趋势。"),
        ],
        [1.35, 2.2, 2.95],
        font_size=8.9,
    )
    add_heading(doc, "4.1 核心业务场景", 2)
    for item in [
        "活动页改版后，增长 PM 需要在实验运行期实时确认 B 组是否提升购买转化率。",
        "投放渠道 GMV 波动时，运营需要区分是渠道质量变化、漏斗异常，还是实验组差异导致。",
        "实验运行到一定样本量后，团队需要基于显著性结果决定继续观察、停止实验或全量上线。",
        "当转化率低于预警阈值时，业务方需要快速定位异常发生在哪个实验组、渠道或漏斗步骤。",
    ]:
        add_number(doc, item)

    add_heading(doc, "5. 指标体系与口径", 1)
    add_body(
        doc,
        "指标口径是本产品的核心资产。所有指标默认基于用户选择的时间窗口计算；涉及人数的指标使用 user_id 去重；涉及订单金额的指标仅统计 buy 事件中的 order_value。"
    )
    add_table(
        doc,
        ["指标", "计算口径", "业务解释", "限制/注意事项"],
        [
            ("DAU", "时间窗口内发生任意事件的去重 user_id 数。", "衡量看板覆盖的活跃用户规模。", "模拟数据不代表真实注册或登录用户。"),
            ("GMV", "buy 事件 order_value 求和。", "衡量交易规模。", "未考虑退款、取消订单和税费。"),
            ("点击人数", "发生 click 的去重 user_id 数。", "漏斗起点样本量。", "同一用户多次点击只计一次。"),
            ("加购人数", "发生 add_to_cart 的去重 user_id 数。", "衡量商品兴趣和购买意向。", "需与点击人数同窗口比较。"),
            ("购买人数", "发生 buy 的去重 user_id 数。", "衡量最终转化用户规模。", "一个用户多次购买在人数上只计一次。"),
            ("加购率", "加购人数 / 点击人数。", "衡量从浏览到意向的转化。", "点击人数为 0 时不展示。"),
            ("购买转化率", "购买人数 / 点击人数。", "A/B 实验的主指标。", "样本不足时不输出确定结论。"),
            ("AOV", "GMV / 订单数。", "衡量平均订单金额。", "订单数为 0 时不展示。"),
            ("uplift", "(B 组指标 - A 组指标) / A 组指标。", "衡量实验组相对提升。", "A 组指标为 0 时不计算。"),
            ("p-value", "two-proportion z-test 检验 A/B 购买转化率差异。", "判断差异是否可能由随机波动造成。", "不等同于业务收益大小，需结合样本量和 uplift。"),
        ],
        [1.05, 2.1, 1.75, 1.6],
        font_size=8.2,
    )

    add_heading(doc, "6. 功能需求", 1)
    add_table(
        doc,
        ["模块", "优先级", "功能描述", "关键产出"],
        [
            ("实时概览大屏", "P0", "展示最近 24 小时或自定义时间范围内的 DAU、GMV、购买转化率、AOV 和趋势图。", "经营概览、趋势判断、预警状态。"),
            ("漏斗分析模块", "P0", "展示 click → add_to_cart → buy 三步漏斗，支持整体和 A/B 组对比。", "各步骤人数、转化率、流失率。"),
            ("A/B 实验评估", "P0", "比较 A 组和 B 组的主指标与辅助指标，自动计算 uplift 和 p-value。", "实验是否显著优于对照组的结论。"),
            ("归因分析模块", "P1", "支持首次触达和最后触达模型切换，展示渠道订单数、GMV 和贡献占比。", "渠道贡献排行、模型切换结果。"),
            ("策略配置模块", "P1", "允许配置实验分流比例、GMV 预警阈值和转化率预警阈值。", "可调整的业务策略参数。"),
        ],
        [1.2, 0.65, 3.0, 1.65],
        font_size=8.4,
    )
    add_heading(doc, "6.1 实时概览大屏", 2)
    for item in [
        "默认展示最近 24 小时整体数据，并支持按小时或按天聚合。",
        "核心卡片包括 DAU、GMV、购买转化率、AOV；低于阈值时展示预警状态。",
        "趋势图至少包含 GMV 趋势和购买转化率趋势，随筛选条件联动刷新。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.2 漏斗分析模块", 2)
    for item in [
        "漏斗步骤固定为 click、add_to_cart、buy，第一版不支持自定义漏斗。",
        "展示每一步去重用户数、相邻步骤转化率和累计转化率。",
        "当出现 buy 人数大于 add_to_cart 人数等模拟数据异常时，后端需返回异常标记供前端提示。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.3 归因分析模块", 2)
    for item in [
        "首次触达模型用于衡量渠道拉新贡献，取用户时间窗口内最早 channel。",
        "最后触达模型用于衡量临门转化贡献，取购买前最近一次非 buy 事件 channel。",
        "归因结果展示渠道 GMV、订单数和贡献占比，支持与时间范围联动。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "6.4 A/B 实验评估模块", 2)
    for item in [
        "A 组定义为对照组，B 组定义为实验组。",
        "主指标为购买转化率，辅助指标包括 GMV、AOV、加购率、订单数。",
        "模块需输出一句业务结论，避免只展示统计值而不支持决策。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. 数据与接口需求", 1)
    add_heading(doc, "7.1 数据字段", 2)
    add_table(
        doc,
        ["数据表", "关键字段", "说明"],
        [
            ("events", "event_id, user_id, session_id, event_type, experiment_group, channel, product_id, order_value, created_at", "核心行为事件表，用于实时指标、漏斗、归因和实验评估。"),
            ("users", "user_id, first_seen_at, acquisition_channel, country, device_type", "用户维表，用于后续人群下钻和渠道分析扩展。"),
            ("experiment_config", "experiment_id, traffic_split_a, traffic_split_b, conversion_alert_threshold, gmv_alert_threshold, updated_at", "实验配置表，用于分流比例和预警阈值管理。"),
        ],
        [1.25, 3.2, 2.05],
        font_size=8.2,
    )
    add_heading(doc, "7.2 API 清单", 2)
    add_table(
        doc,
        ["接口", "方法", "入参", "返回结果"],
        [
            ("/health", "GET", "无", "服务状态、数据库连接状态。"),
            ("/api/metrics", "GET", "start_time, end_time, granularity, experiment_group", "DAU、GMV、转化率、AOV、趋势数据。"),
            ("/api/funnel", "GET", "start_time, end_time, experiment_group", "漏斗各步骤人数、转化率、流失率。"),
            ("/api/attribution", "GET", "start_time, end_time, model", "渠道订单数、GMV、贡献占比。"),
            ("/api/experiment", "GET", "start_time, end_time", "A/B 指标对比、uplift、p-value、实验结论。"),
            ("/api/config", "GET/POST", "traffic_split, alert_thresholds", "读取或更新实验配置与预警阈值。"),
        ],
        [1.25, 0.65, 2.35, 2.25],
        font_size=8.1,
    )

    add_heading(doc, "8. 实验评估规则", 1)
    add_body(
        doc,
        "实验评估模块的产品目标是把统计检验结果翻译成业务可读的行动建议。第一版使用购买转化率作为主指标，以 A 组为对照组，B 组为实验组。"
    )
    add_table(
        doc,
        ["结论类型", "触发条件", "前端展示建议"],
        [
            ("样本量不足", "A 组或 B 组点击人数低于最小样本阈值。", "提示继续观察，不给出上线建议。"),
            ("显著优于", "p-value < 0.05 且 B 组购买转化率高于 A 组。", "提示实验组显著优于对照组，可考虑扩大流量或全量上线。"),
            ("无显著差异", "p-value >= 0.05。", "提示当前数据不足以证明实验组优于对照组，建议继续观察或复盘策略。"),
            ("显著低于", "p-value < 0.05 且 B 组购买转化率低于 A 组。", "提示实验组存在负向影响，应暂停实验并排查原因。"),
        ],
        [1.2, 2.65, 2.65],
        font_size=8.5,
    )
    add_callout(
        doc,
        "PM 解释口径",
        "p-value 只回答“差异是否可能来自随机波动”，不直接回答“业务收益是否足够大”。最终上线判断需要同时看显著性、uplift、GMV 影响、样本量和策略风险。",
    )

    add_heading(doc, "9. 交互与页面说明", 1)
    for item in [
        "用户进入页面后，默认看到最近 24 小时整体经营概览。",
        "顶部筛选区包含时间范围、时间粒度、实验组、归因模型和自动刷新开关。",
        "切换实验组后，概览卡片、趋势图、漏斗图和实验对比模块同步刷新。",
        "切换归因模型后，渠道贡献图重新计算，页面保留当前时间范围。",
        "策略配置保存后，预警阈值立即影响概览卡片的状态展示。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "9.1 页面信息架构", 2)
    add_table(
        doc,
        ["区域", "内容", "设计重点"],
        [
            ("顶部筛选区", "时间范围、粒度、实验组、归因模型、自动刷新。", "所有分析模块共用筛选条件，减少理解成本。"),
            ("概览区", "DAU、GMV、转化率、AOV 卡片与趋势图。", "先回答业务是否健康。"),
            ("诊断区", "漏斗分析、渠道归因。", "再回答哪里出现变化。"),
            ("决策区", "A/B 实验评估、自动结论、策略配置。", "最后支持继续实验、停止实验或上线决策。"),
        ],
        [1.2, 2.65, 2.65],
        font_size=8.7,
    )

    add_heading(doc, "10. 验收标准", 1)
    add_table(
        doc,
        ["验收类型", "验收标准"],
        [
            ("本地运行", "FastAPI、Streamlit、数据模拟器均可通过 README 指令启动。"),
            ("数据模拟", "模拟器能持续写入 click、add_to_cart、buy 事件，并包含 user_id、experiment_group、channel、created_at。"),
            ("API 指标", "/api/metrics、/api/funnel、/api/experiment 能返回非空结果，字段命名稳定。"),
            ("漏斗逻辑", "正常模拟数据下漏斗人数满足 click ≥ add_to_cart ≥ buy。"),
            ("实验评估", "当 B 组转化率更高且样本充足时，系统能返回 uplift、p-value 和明确结论。"),
            ("前端体验", "筛选条件改变后，图表联动刷新；图表无明显遮挡或空白异常。"),
            ("部署上线", "Render 后端健康检查通过，Streamlit Cloud 可访问并能请求后端 API。"),
        ],
        [1.45, 5.05],
        font_size=8.8,
    )

    add_heading(doc, "11. 风险与后续迭代", 1)
    add_heading(doc, "11.1 主要风险", 2)
    add_table(
        doc,
        ["风险", "影响", "应对策略"],
        [
            ("样本量不足", "实验结论不稳定，容易误判。", "设置最小样本阈值，样本不足时只提示观察。"),
            ("模拟数据偏离真实业务", "作品集 Demo 的数据分布可能过于理想。", "在 README 和 PRD 中说明模拟假设，后续支持真实埋点接入。"),
            ("归因模型简化", "首次/最后触达无法覆盖多触点贡献。", "作为第一版产品边界，后续扩展线性归因、时间衰减归因。"),
            ("实时性有限", "轮询刷新无法达到真正流计算秒级能力。", "第一版采用短间隔刷新，后续引入流处理或缓存层。"),
            ("配置误操作", "错误阈值或分流比例可能影响判断。", "后续增加权限、变更记录和配置回滚。"),
        ],
        [1.4, 2.2, 2.9],
        font_size=8.7,
    )
    add_heading(doc, "11.2 后续迭代方向", 2)
    for item in [
        "多实验管理：支持多个实验并行、实验状态流转和实验历史沉淀。",
        "人群下钻：支持新老用户、设备、国家、渠道等维度分析。",
        "高级归因：扩展线性归因、时间衰减归因和 Shapley Value 归因。",
        "实验方法升级：增加贝叶斯实验评估、序贯检验和最小可检测效应配置。",
        "数据治理：补充指标血缘、口径版本、异常事件校验和埋点质量监控。",
        "协作能力：增加结论导出、实验复盘模板和决策记录。"
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. 作品集表达重点", 1)
    add_body(
        doc,
        "面试展示时，该项目应强调你不是只做了一个 Dashboard，而是完成了从业务问题抽象、指标口径定义、实验判断规则、数据接口设计到可视化决策闭环的一整套数据产品设计。"
    )
    add_table(
        doc,
        ["能力维度", "项目体现"],
        [
            ("业务理解", "围绕综合电商转化增长和投放优化构建场景。"),
            ("指标设计", "定义 DAU、GMV、转化率、AOV、uplift、p-value 等口径。"),
            ("实验思维", "用显著性检验和样本量判断支持上线决策。"),
            ("归因分析", "区分首次触达和最后触达模型的业务适用场景。"),
            ("产品落地", "把数据模拟、后端 API、前端看板和部署路径串成完整闭环。"),
        ],
        [1.35, 5.15],
        font_size=8.9,
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
